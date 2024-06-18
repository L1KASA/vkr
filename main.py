from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx import Document
import matplotlib
import matplotlib.pyplot as plt
import numpy as np
import flet as ft
from flet.matplotlib_chart import MatplotlibChart
import os
import re
import numpy as np
import matplotlib.pyplot as plt
from deap import base, creator, tools, algorithms
import optuna

matplotlib.use("svg")
temp_filenames = ["temp_plot1.png", "temp_plot2.png", "temp_plot3.png", "temp_plot4.png"]

def euler_difference_equations(a, b, c, rx, ry, g, f, k, q, X0, Y0, time_points):
    """
    Реализация метода Эйлера для решения разностных уравнений с заданными временными точками.

    Args:
    - a, b, c, rx, ry, g, f, k, q: параметры уравнения
    - X0, Y0: начальные значения популяций
    - time_points: поколения

    Returns:
    - список значений X и Y по заданным временным точкам
    """
    results_X = [X0]
    results_Y = [Y0]
    
    current_time = time_points[0]
    
    for i in range(1, len(time_points)):
        next_time_point = time_points[i]
        while current_time < next_time_point:
            x = results_X[-1] / (results_X[-1] + results_Y[-1])
            y = results_Y[-1] / (results_X[-1] + results_Y[-1])

            term1_X = (rx * a * x + ry * b * y) * results_X[-1]
            term2_X = (rx * b * x + ry * c * y) * results_Y[-1]
            term3_X = g * results_X[-1]
            term4_X = f * results_X[-1] * results_X[-1]
            term5_X = k * results_X[-1] * results_Y[-1]

            term1_Y = (rx * (1 - a) * x + ry * (1 - b) * y) * results_X[-1]
            term2_Y = (rx * (1 - b) * x + ry * (1 - c) * y) * results_Y[-1]
            term3_Y = q * results_Y[-1]

            h = next_time_point - current_time

            next_X = results_X[-1] + h * (term1_X + term2_X - term3_X - term4_X - term5_X)
            next_Y = results_Y[-1] + h * (term1_Y + term2_Y - term3_Y)

            current_time += h

        results_X.append(next_X)
        results_Y.append(next_Y)

    return results_X, results_Y

# Функция Гаусса
def gaussian(x, a, b, c):
    return a * np.exp(-((x - b) ** 2) / (2 * c ** 2))

# Вычисление градиента функции ошибок
def gradient(params, x_data, y_data):
    a, b, c = params
    predictions = gaussian(x_data, a, b, c)
    errors = y_data - predictions
    grad_a = -2 * np.sum(errors * np.exp(-((x_data - b) ** 2) / (2 * c ** 2)))
    grad_b = -2 * np.sum(errors * a * (x_data - b) / c**2 * np.exp(-((x_data - b) ** 2) / (2 * c ** 2)))
    grad_c = -2 * np.sum(errors * a * (x_data - b) ** 2 / c**3 * np.exp(-((x_data - b) ** 2) / (2 * c ** 2)))
    return np.array([grad_a, grad_b, grad_c])

# Функция стоимости (сумма квадратов ошибок)
def cost_function(params, x_data, y_data):
    predictions = gaussian(x_data, params[0], params[1], params[2])
    errors = y_data - predictions
    return np.sum(errors ** 2)

def mgp(x_data, y_data, learning_rate=0.001, num_iterations=10000, tolerance=1e-6):
    params = np.array([max(y_data), np.mean(x_data), np.std(x_data)])   
    costs = []  # Список для хранения значений функции стоимости на каждой итерации
    prev_cost = float('inf')  # Инициализация предыдущей стоимости большим значением

    for iteration in range(num_iterations):
        grad = gradient(params, x_data, y_data)
        params -= learning_rate * grad
        cost = cost_function(params, x_data, y_data)
        costs.append(cost)
        
        # Условие остановки по изменению значения функции стоимости
        if abs(prev_cost - cost) < tolerance:
            print(f"Остановлено на {iteration+1}-й итерации по изменению стоимости.")
            break
        # Условие остановки по изменению параметров
        if np.linalg.norm(learning_rate * grad) < tolerance:
            print(f"Остановлено на {iteration+1}-й итерации по изменению параметров.")
            break

        prev_cost = cost  # Обновление предыдущей стоимости

    return params, costs

def main(page: ft.Page):
    page.window_height = 700
    page.window_width = 950
    page.update()
    
    def window_event(e):
        if e.data == "close":
            print('Приложение закрыто.')
        
            for filename in temp_filenames:
                if os.path.exists(filename):
                    os.remove(filename) 
            page.window_destroy()

    page.window_prevent_close = True
    page.on_window_event = window_event
    page.scroll = "AUTO"   
    page.theme_mode = ft.ThemeMode.LIGHT
    
    def switch_data(e):
        if switch_real_or_gauss.value:
            switch_real_or_gauss.label = "Реальные данные"
        else:
            switch_real_or_gauss.label = "Данные с функцией Гаусса"
        page.update()

    def select_auto_or_not(e):
        if switch_autoselect_or_not.value:
            switch_autoselect_or_not.label = "Автоматическая настройка генетического алгоритма"

            POPULATION_SIZE.disabled = True
            P_CROSSOVER.disabled = True
            P_MUTATION.disabled = True
            TOURNSIZE.disabled = True
            MAX_GENERATIONS.disabled = True
            MU.disabled = True
            SIGMA.disabled = True
            INDPB.disabled = True
        else:
            switch_autoselect_or_not.label = "Настройка гиперпараметров вручную"

            POPULATION_SIZE.disabled = False
            P_CROSSOVER.disabled = False
            P_MUTATION.disabled = False
            TOURNSIZE.disabled = False
            MAX_GENERATIONS.disabled = False
            MU.disabled = False
            SIGMA.disabled = False
            INDPB.disabled = False
        page.update()
    
    switch_real_or_gauss = ft.Switch(label="Данные с функцией Гаусса", on_change=switch_data)
    switch_autoselect_or_not = ft.Switch(label="Настройка гиперпараметров вручную", on_change=select_auto_or_not)

    def fitness_function(individual):
        a, b, c, rx, ry, g, f, k, q = individual

        X_model, Y_model = euler_difference_equations(a, b, c, rx, ry, g, f, k, q, 1, 1, years)
        
        # Нормализация значений
        X_model_max = max(X_model)
        X_model_normalized = [x / X_model_max for x in X_model]

        difference = None
        if switch_real_or_gauss.value:
            difference = np.array(areas) - np.array(X_model_normalized)

        else:
           difference = np.array(y_values) - np.array(X_model_normalized)

        # Евклидова норма разности
        norm_difference = np.sqrt(np.sum(difference ** 2))
 
        # Проверка на бесконечность (inf) или не число (NaN)
        if np.isinf(norm_difference) or np.isnan(norm_difference):
            # Замена inf или NaN на значение по умолчанию
            default_value = 10000 
            return (default_value,)
        return norm_difference,

    def theme_changed(e):
        page.theme_mode = (
            ft.ThemeMode.DARK
            if page.theme_mode == ft.ThemeMode.LIGHT
            else ft.ThemeMode.LIGHT
        )
        c.label = (
            "Светлая тема" if page.theme_mode == ft.ThemeMode.LIGHT else "Темная тема"
        )
        page.update()

    page.theme_mode = ft.ThemeMode.LIGHT
    c = ft.Switch(label="Светлая тема", on_change=theme_changed)
    
    dlg = ft.AlertDialog(
        title=ft.Text(\
            "В первую строку необходимо вписать год, во вторую площадь территории. Между значениями может быть пробел, запятая и точка с запятой. Перед каждой датой до нашей эры прописывается знак минуса.")\
            , on_dismiss=lambda e: print("Диалоговое окно закрыто.")
    )

    dlg1 = ft.AlertDialog(
        title=ft.Text(\
            "Тема: Исследование природно-климатических факторов на исторические процессы методом математического моделирования.\nПочта автора: tahmazova_anzhelika@mail.ru")\
            , on_dismiss=lambda e: print("Диалоговое окно закрыто.")
    )

    def handle_menu_item_click(e):
        page.dialog = dlg
        dlg.open = True
        page.update()

    def handle_menu_item_click1(e):
        page.dialog = dlg1
        dlg1.open = True
        page.update()

    def handle_on_open(e):
        print(f"{e.control.content.value}.on_open")

    def handle_on_close(e):
        print(f"{e.control.content.value}.on_close")

    def handle_on_hover(e):
        print(f"{e.control.content.value}.on_hover")

    def exit_fun(e):
        for filename in temp_filenames:
            if os.path.exists(filename):
                print(f"Файл {filename} существует.")
                os.remove(filename)
            else:
                print(f"Файл {filename} не существует.")
        page.window_destroy()
    
    menubar = ft.MenuBar(
        expand=True,
        style=ft.MenuStyle(
            alignment=ft.alignment.top_left,
            bgcolor= ft.colors.BLUE_400,
            mouse_cursor={ft.MaterialState.HOVERED: ft.MouseCursor.WAIT,
                          ft.MaterialState.DEFAULT: ft.MouseCursor.ZOOM_OUT},
        ),
        controls=[
            ft.SubmenuButton(
                content=ft.Text("Файл"),
                on_open=handle_on_open,
                on_close=handle_on_close,
                on_hover=handle_on_hover,
                controls=[
                    ft.MenuItemButton(
                        content=ft.Text("О программе"),
                        leading=ft.Icon(ft.icons.INFO),
                        style=ft.ButtonStyle(bgcolor={ft.MaterialState.HOVERED: ft.colors.GREEN_500}),
                        on_click=handle_menu_item_click1
                    ),
                    ft.MenuItemButton(
                        content=ft.Text("Сохранить"),
                        leading=ft.Icon(ft.icons.SAVE),
                        style=ft.ButtonStyle(bgcolor={ft.MaterialState.HOVERED: ft.colors.GREEN_500}),
                        on_click=lambda _: saveme.save_file()
                    ),
                    ft.MenuItemButton(
                        content=ft.Text("Закрыть приложение"),
                        leading=ft.Icon(ft.icons.CLOSE),
                        style=ft.ButtonStyle(bgcolor={ft.MaterialState.HOVERED: ft.colors.GREEN_500}),
                        on_click=exit_fun
                    )
                ]
            ),
            ft.MenuItemButton(
                        content=ft.Text("Инструкция по работе с программой"),                        
                        on_click=handle_menu_item_click
                    )
        ]
    )

    def tb3_changes(field, pagee):
        c1.disabled = False
        c2.disabled = False
        c1.update()
        c2.update()
        global last_valid_value
        digits = field.value
        if digits == "":
            last_valid_value = ""
        else:
            print(field.value)

    tb3 = ft.TextField(label="Год. Пример: 1210, 1225, 1300", hint_text="1210, 1225, 1300",  
                       on_change = lambda e: tb3_changes(tb3, page))

    tb4 = ft.TextField(label="Площадь территории. Пример: 1.0, 2.4, 1.6", hint_text="1.0, 2.4, 1.6",  
                       on_change = lambda e: tb3_changes(tb4, page))
    
    # количество индивидуумов в популяции
    POPULATION_SIZE = ft.TextField(label="Количество индивидуумов в популяции", hint_text="250", value=250, 
                                   on_change = lambda e: tb3_changes(POPULATION_SIZE, page))
    # вероятность скрещивания
    P_CROSSOVER = ft.TextField(label="Вероятность скрещивания", hint_text="0.8", value=0.8, 
                               on_change = lambda e: tb3_changes(P_CROSSOVER, page))
    # вероятность мутации индивидуума
    P_MUTATION = ft.TextField(label="Вероятность мутации", hint_text="0.2", value=0.2, 
                              on_change = lambda e: tb3_changes(P_MUTATION, page))
    TOURNSIZE = ft.TextField(label="Размер турнира", hint_text="3",  value=3, 
                             on_change = lambda e: tb3_changes(TOURNSIZE, page))
    # максимальное количество поколений
    MAX_GENERATIONS = ft.TextField(label="Максимальное количество поколений", hint_text="50", value=50, 
                                   on_change = lambda e: tb3_changes(MAX_GENERATIONS, page))
    MU = ft.TextField(label="Математическое ожидание", 
                      on_change = lambda e: tb3_changes(MU, page))
    SIGMA = ft.TextField(label="Стандартное отклонение", 
                                   on_change = lambda e: tb3_changes(SIGMA, page))
    INDPB = ft.TextField(label="Вероятность мутации гена в индивидууме", 
                                   on_change = lambda e: tb3_changes(INDPB, page))

    def c1_click(d):
        print(tb4.value)
        if str(tb3.value) == '' or str(tb4.value) == '':
            print('Ошибка! Данные отсутствуют.')
            page.show_snack_bar(ft.SnackBar(content=ft.Text("Ошибка! Данные отсутствуют.")))
            return
        
        if len(np.array(re.split('[,;]\s|,|;|\s', tb3.value)).astype(int)) != len(np.array(re.split('[,;]\s|,|;|\s', tb4.value)).astype(float)):
            print('Ошибка! Количество данных в строках не совпадает.')
        
            page.show_snack_bar(ft.SnackBar(content=ft.Text('Ошибка! Количество данных в строках не совпадает.')))
            return
        
        global years_no_normalized
        global areas_no_normalized

        years_no_normalized = np.array(re.split('[,;]\s|,|;|\s', tb3.value)).astype(int)
        areas_no_normalized = np.array(re.split('[,;]\s|,|;|\s', tb4.value)).astype(float)

        global years
        years = (years_no_normalized - years_no_normalized.min()) / 25
       
        global areas
        areas = (areas_no_normalized - areas_no_normalized.min()) / (areas_no_normalized.max() - areas_no_normalized.min())

        # Генерация значений функции Гаусса с оптимальными параметрами
        global y_values

        optimal_params, costs = mgp(years, areas)
        print(costs[-1])
        y_values = gaussian(years, *optimal_params)
        max_y_value = max(y_values)
        y_values /= max_y_value

        creator.create("FitnessMin", base.Fitness, weights=(-1.0,))
        creator.create("Individual", list, fitness=creator.FitnessMin)
        
        def create_toolbox(toolbox):
            toolbox.register("attr_float", np.random.uniform, 0, 1)
            toolbox.register("individual", tools.initRepeat, creator.Individual, toolbox.attr_float, n=9)
            toolbox.register("population", tools.initRepeat, list, toolbox.individual)
            toolbox.register("evaluate", fitness_function)  
        global best_individual
        global best_individual_fitness
        if switch_autoselect_or_not.value:
            # Различные операторы для вещественного кодирования
            crossover_ops = {
                'cxOnePoint': tools.cxOnePoint,
                'cxTwoPoint': tools.cxTwoPoint,
                #'cxBlend': tools.cxBlend
            }

            mutation_ops = {
                'mutGaussian': tools.mutGaussian,
            }

            selection_ops = {
                'selTournament': lambda ind, k, tournsize: tools.selTournament(ind, k, tournsize=tournsize),
                'selRoulette': tools.selRoulette,
                'selRank': tools.selStochasticUniversalSampling,
                'selBest': tools.selBest
            }
            
            # Функция для оценки эффективности гиперпараметров
            def evaluate_params(trial):
                crossover = trial.suggest_categorical('crossover', list(crossover_ops.keys()))
                mutation = trial.suggest_categorical('mutation', list(mutation_ops.keys()))
                selection = trial.suggest_categorical('selection', list(selection_ops.keys()))

                mu = trial.suggest_float('mu', 0, 1)
                sigma = trial.suggest_float('sigma', 0, 1)
                indpb = trial.suggest_float('indpb', 0.005, 0.01)
                tournsize = trial.suggest_int('tournsize', 2, 10)
                population_size = trial.suggest_int('population_size', 40, 250)
                cxpb = trial.suggest_float('cxpb', 0.6, 1.0)
                mutpb = trial.suggest_float('mutpb', 0.005, 0.01)
                ngen = trial.suggest_int('ngen', 10, 70)

                # Создаем новый инстанс Toolbox
                toolbox = base.Toolbox()
                create_toolbox(toolbox)
                # Регистрация операторов кроссинговера
                toolbox.register("mate", crossover_ops[crossover])

                # Регистрация операторов мутации
                if mutation == 'mutGaussian':
                    toolbox.register("mutate", mutation_ops[mutation], mu=mu, sigma=sigma, indpb=indpb)

                # Регистрация операторов селекции
                if selection == 'selTournament':
                    toolbox.register("select", selection_ops[selection], tournsize=tournsize)
                else:
                    toolbox.register("select", selection_ops[selection])
                
                population = toolbox.population(n=population_size)
                algorithms.eaSimple(population, toolbox, cxpb, mutpb, ngen)    

                best_individual = tools.selBest(population, 1)[0]
                return best_individual.fitness.values[0]

            # Оптимизация гиперпараметров с использованием Optuna
            study = optuna.create_study(direction='minimize')
            study.optimize(evaluate_params, n_trials=50)
            global best_params
            best_params = study.best_params
            print(f"Лучшие гиперпараметры: {best_params}")
            print(f"Лучший результат: {study.best_value}")
           
            toolbox = base.Toolbox()
            create_toolbox(toolbox)

            # Регистрация операторов кроссинговера
            toolbox.register("mate", crossover_ops[study.best_params['crossover']])

            # Регистрация операторов мутации
            if study.best_params['mutation'] == 'mutGaussian':
                toolbox.register("mutate", mutation_ops[study.best_params['mutation']], mu=study.best_params['mu'], sigma=study.best_params['sigma'], indpb=study.best_params['indpb'])

            # Регистрация операторов селекции
            if study.best_params['selection'] == 'selTournament':
                toolbox.register("select", selection_ops[study.best_params['selection']], tournsize=study.best_params['tournsize'])
            else:
                toolbox.register("select", selection_ops[study.best_params['selection']])

            statistics = tools.Statistics(lambda ind: ind.fitness.values)
            statistics.register("min", np.min)
            
            population = tools.initRepeat(list, toolbox.individual, n=study.best_params['population_size'])

            population, logbook = algorithms.eaSimple(
                population,
                toolbox,
                cxpb=study.best_params['cxpb'],
                mutpb=study.best_params['mutpb'],
                ngen=study.best_params['ngen'],
                stats=statistics,
            )

            min_fitness_values = logbook.select("min")
            
            best_individual = tools.selBest(population, k=1)[0]
            best_individual_fitness = fitness_function(best_individual)
            page.update()
        else:
            toolbox = base.Toolbox()
            create_toolbox(toolbox)

            toolbox.register("mate", tools.cxTwoPoint)

            toolbox.register("mutate", tools.mutGaussian, mu=float(MU.value), sigma= float(SIGMA.value), indpb=float(INDPB.value))

            toolbox.register("select", tools.selTournament, tournsize=int(TOURNSIZE.value))
            
            page.update()
            statistics = tools.Statistics(lambda ind: ind.fitness.values)
            statistics.register("min", np.min)

            population = toolbox.population(n=int(POPULATION_SIZE.value))
            population, logbook = algorithms.eaSimple(population, toolbox, cxpb=float(P_CROSSOVER.value), mutpb=float(P_MUTATION.value), 
                                ngen=int(MAX_GENERATIONS.value), stats=statistics, verbose=True)

            min_fitness_values = logbook.select("min")
            
            best_individual = tools.selBest(population, k=1)[0]
            best_individual_fitness = fitness_function(best_individual)
            page.update()

        a, b, c, rx, ry, g, f, k, q = best_individual
        
        # Решение уравнений методом Эйлера
        X_model, Y_model = euler_difference_equations(a, b, c, rx, ry, g, f, k, q, 1, 1, years)
        # Нормализация значений
        X_model_max = max(X_model)
        global X_model_normalized
        X_model_normalized = [x / X_model_max for x in X_model]

        # Построение графиков 
        fig1, ax1 = plt.subplots()  
        ax1.plot(years, y_values, label='Функция Гаусса', color='blue')
        ax1.scatter(years, areas, label='Реальные данные', color='red')
        ax1.set_xlabel("Поколения")
        ax1.set_ylabel("Пассионарность")
        ax1.set_title('Гауссова функция')
        ax1.legend()
        ax1.grid(True)

        fig2, ax2 = plt.subplots() 
        ax2.plot(years, y_values, label='Функция Гаусса', color='blue')
        ax2.scatter(years, areas, label='Реальные данные', color='red') 
        ax2.plot(years, X_model_normalized, label='Пассионарии (Модель)', color='green')
        ax2.set_xlabel("Поколения")
        ax2.set_ylabel("Пассионарность")
        ax2.set_title('Гауссова функция и модель этногенеза')
        ax2.legend()
        ax2.grid(True)

        fig3, ax3 = plt.subplots() 
        ax3.plot(min_fitness_values, color='blue')
        ax3.set_xlabel("Поколения")
        ax3.set_ylabel("Функция пригодности")
        ax3.set_title('Зависимость значений функции пригодности от поколения')
        ax3.legend()
        ax3.grid(True)

        fig4, ax4 = plt.subplots() 
        ax4.scatter(years, areas, label='Реальные данные', color='red') 
        ax4.plot(years, X_model_normalized, label='Пассионарии (Модель)', color='green')
        ax4.set_xlabel("Поколения")
        ax4.set_ylabel("Пассионарность")
        ax4.set_title('Математическая модель этногенеза')
        ax4.legend()
        ax4.grid(True)

        fig1.savefig(temp_filenames[0], format="png", bbox_inches="tight", pad_inches=0.1)
        fig2.savefig(temp_filenames[1], format="png", bbox_inches="tight", pad_inches=0.1)
        fig3.savefig(temp_filenames[2], format="png", bbox_inches="tight", pad_inches=0.1)
        fig4.savefig(temp_filenames[3], format="png", bbox_inches="tight", pad_inches=0.1)

        page.update()

        chart1.figure = fig1
        chart2.figure = fig2
        chart3.figure = fig3

        chart1.update()
        chart2.update()
        chart3.update()
    
    def c2_click(d):
        tb3.value = tb4.value = ''
        page.update()

    c1 = ft.Container(
        content=ft.ElevatedButton("Вычислить"),
        padding=5,
        on_click=c1_click,
        disabled=True       
    )
    c2 = ft.Container(
        content=ft.ElevatedButton("Очистить"),
        padding=5,
        on_click=c2_click,
        disabled=True       
    )

    def mysavefile(e:ft.FilePickerResultEvent):
        for temp_filename in temp_filenames:
            if not os.path.exists(temp_filename):
                print(f"Отсутствует путь к файлу '{temp_filename}'.")
                page.show_snack_bar(ft.SnackBar(content=ft.Text("Ошибка! Графики не были нарисованы.")))
                return
        save_loc = e.path
        if not save_loc:
            print("Ошибка при сохранении файла", e)
            page.show_snack_bar(ft.SnackBar(content=ft.Text("Ошибка! Документ не сохранился.")))
            return
        
        document = Document()
        style = document.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)

        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        document.add_paragraph('Для вычисления евклидовой метрики в генетическом алгоритме используются: ' 
        + str(switch_real_or_gauss.label).lower())

        document.add_paragraph('')
        document.add_picture(temp_filenames[0])
        style = document.paragraphs[-1]
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph("Рисунок 1 – Исходный график с данными")
        style = document.paragraphs[-1]
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_picture(temp_filenames[1])

        style = document.paragraphs[-1]
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph("Рисунок 2 – Исходный график с данными")

        style = document.paragraphs[-1]
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_picture(temp_filenames[2])

        style = document.paragraphs[-1]
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph("Рисунок 3 – Исходный график с данными")
        style = document.paragraphs[-1]

        style.alignment = 1

        style = document.paragraphs[-1]
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_picture(temp_filenames[3])

        style = document.paragraphs[-1]
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph("Рисунок 4 – Исходный график с данными")
        style = document.paragraphs[-1]

        style.alignment = 1
        
        document.add_paragraph('')
        document.add_paragraph('Поколение: площадь территории')
        for year, area in dict(zip(years_no_normalized, areas_no_normalized)).items():
            para = document.add_paragraph("")
            para.add_run(str(year) + ': ' + str(area))
        document.add_paragraph('Гиперпараметры: ')
        if switch_autoselect_or_not.value:
            document.add_paragraph('Лучшие гипараметры и операторы: ' + str(best_params))
            
        else:
            document.add_paragraph('MU: ' + str(MU.value))
            document.add_paragraph('SIGMA: ' + str(SIGMA.value))
            document.add_paragraph('INDPB: ' + str(INDPB.value))
            document.add_paragraph('POPULATION_SIZE: ' + str(POPULATION_SIZE.value))
            document.add_paragraph('POPULATION_SIZE: ' + str(POPULATION_SIZE.value))
            document.add_paragraph('P_CROSSOVER: ' + str(P_CROSSOVER.value))
            document.add_paragraph('POPULATION_SIZE: ' + str(POPULATION_SIZE.value))
            document.add_paragraph('P_MUTATION: ' + str(P_MUTATION.value))
            document.add_paragraph('TOURNSIZE: ' + str(TOURNSIZE.value))
            document.add_paragraph('MAX_GENERATIONS: ' + str(MAX_GENERATIONS.value))

        document.add_paragraph('')
        document.add_paragraph('Поколение: пассионарность из математической модели')
        for i in range(len(X_model_normalized)):
            para = document.add_paragraph("")
            para.add_run(str(years[i]) + ': ' + str(X_model_normalized[i]))
        document.add_paragraph("Лучший индивид: " + str(round(float(*best_individual_fitness), 3)))

        document.add_paragraph('Найденные параметры: ')

        for key, value in dict(zip(['a', 'b', 'c', 'rx', 'ry', 'g', 'f', 'k', 'q'], best_individual)).items():
            para = document.add_paragraph("")
            para.add_run(key + ' = ' + str(round(value, 3))) 
        
        if '.docx' not in save_loc:
            save_loc = save_loc + '.docx'
        document.save(save_loc)
        
        page.show_snack_bar(ft.SnackBar(content=ft.Text("Графики и данные успешно сохранены в документе Word")))

        page.update()

    saveme = ft.FilePicker(on_result=mysavefile)
    page.overlay.append(saveme)

    fig, ax = plt.subplots()
    chart1 = MatplotlibChart(fig, isolated=True, expand=True)
    chart2 = MatplotlibChart(fig, isolated=True, expand=True)
    chart3 = MatplotlibChart(fig, isolated=True, expand=True)
    
    page.add(ft.Row([menubar]), tb3, tb4, ft.Row(controls=[c1, c2, switch_real_or_gauss, switch_autoselect_or_not]), 
             ft.Row([P_CROSSOVER, POPULATION_SIZE, P_MUTATION]),ft.Row([TOURNSIZE, MAX_GENERATIONS]),  
             ft.Row([MU, SIGMA, INDPB]), 
             ft.Row([chart1, chart2, chart3]), c)
    
ft.app(target=main)
